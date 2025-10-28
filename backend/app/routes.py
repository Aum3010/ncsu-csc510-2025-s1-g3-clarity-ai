import os
from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
import pypdf
import docx
import json
from datetime import datetime

from .main import db
from .models import Document, Requirement, Tag, ProjectSummary, UserProfile
from .rag_service import (
    process_and_store_document,
    generate_project_requirements,
    generate_project_summary,
    delete_document_from_rag
)
from .auth_service import get_roles_permissions_config, require_auth

api_bp = Blueprint('api', __name__, url_prefix='/api')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md', 'json'}


@api_bp.route('/')
def index():
    return jsonify({"message": "Welcome to the Clarity AI API!"})

# --- Health Check Endpoints ---


@api_bp.route('/health/supertokens', methods=['GET'])
def supertokens_health_check():
    """
    Health check endpoint to verify SuperTokens connectivity and configuration.
    """
    try:
        # Test SuperTokens connectivity by making a simple request to SuperTokens Core
        import requests
        from flask import current_app

        # Get SuperTokens connection URI from config
        supertokens_uri = current_app.config.get(
            'connection_uri', 'http://localhost:3567')

        try:
            # Test connectivity to SuperTokens Core
            response = requests.get(f"{supertokens_uri}/hello", timeout=5)

            if response.status_code == 200:
                return jsonify({
                    "status": "healthy",
                    "service": "SuperTokens",
                    "message": "SuperTokens Core is connected and responding",
                    "core_url": supertokens_uri,
                    "timestamp": datetime.utcnow().isoformat()
                })
            else:
                return jsonify({
                    "status": "unhealthy",
                    "service": "SuperTokens",
                    "message": f"SuperTokens Core responded with status {response.status_code}",
                    "core_url": supertokens_uri,
                    "timestamp": datetime.utcnow().isoformat()
                }), 503

        except requests.exceptions.RequestException as e:
            return jsonify({
                "status": "unhealthy",
                "service": "SuperTokens",
                "message": f"SuperTokens Core connection failed: {str(e)}",
                "core_url": supertokens_uri,
                "timestamp": datetime.utcnow().isoformat()
            }), 503

    except Exception as e:
        return jsonify({
            "status": "error",
            "service": "SuperTokens",
            "message": f"Health check failed: {str(e)}",
            "timestamp": datetime.utcnow().isoformat()
        }), 500


@api_bp.route('/health/database', methods=['GET'])
def database_health_check():
    """
    Health check endpoint to verify database connectivity.
    """
    try:
        # Test database connectivity by running a simple query
        from .main import db
        result = db.session.execute(db.text('SELECT 1')).fetchone()

        if result:
            return jsonify({
                "status": "healthy",
                "service": "Database",
                "message": "Database connection is working",
                "timestamp": datetime.utcnow().isoformat()
            })
        else:
            return jsonify({
                "status": "unhealthy",
                "service": "Database",
                "message": "Database query returned no result",
                "timestamp": datetime.utcnow().isoformat()
            }), 503

    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "service": "Database",
            "message": f"Database connectivity error: {str(e)}",
            "timestamp": datetime.utcnow().isoformat()
        }), 503


@api_bp.route('/health/full', methods=['GET'])
def full_health_check():
    """
    Comprehensive health check endpoint that verifies all system components.
    """
    from flask import current_app

    health_status = {
        "overall_status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "services": {}
    }

    # Check API health
    health_status["services"]["api"] = {
        "status": "healthy",
        "message": "API is running"
    }

    # Check database health
    try:
        from .main import db
        result = db.session.execute(db.text('SELECT 1')).fetchone()
        if result:
            health_status["services"]["database"] = {
                "status": "healthy",
                "message": "Database connection is working"
            }
        else:
            health_status["services"]["database"] = {
                "status": "unhealthy",
                "message": "Database query returned no result"
            }
            health_status["overall_status"] = "degraded"
    except Exception as e:
        health_status["services"]["database"] = {
            "status": "unhealthy",
            "message": f"Database connectivity error: {str(e)}"
        }
        health_status["overall_status"] = "unhealthy"

    # Check SuperTokens health
    try:
        import requests

        # Get SuperTokens connection URI from config
        supertokens_uri = current_app.config.get(
            'connection_uri', 'http://localhost:3567')

        # Test connectivity to SuperTokens Core
        response = requests.get(f"{supertokens_uri}/hello", timeout=5)

        if response.status_code == 200:
            health_status["services"]["supertokens"] = {
                "status": "healthy",
                "message": "SuperTokens Core is connected and responding",
                "core_url": supertokens_uri,
                "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
                "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
            }
        else:
            health_status["services"]["supertokens"] = {
                "status": "unhealthy",
                "message": f"SuperTokens Core responded with status {response.status_code}",
                "core_url": supertokens_uri,
                "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
                "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
            }
            health_status["overall_status"] = "degraded"

    except requests.exceptions.RequestException as e:
        health_status["services"]["supertokens"] = {
            "status": "unhealthy",
            "message": f"SuperTokens Core connection failed: {str(e)}",
            "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
            "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
        }
        health_status["overall_status"] = "degraded"
    except Exception as e:
        health_status["services"]["supertokens"] = {
            "status": "unhealthy",
            "message": f"SuperTokens health check failed: {str(e)}",
            "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
            "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
        }
        health_status["overall_status"] = "degraded"

    # Determine overall status code
    if health_status["overall_status"] == "healthy":
        return jsonify(health_status)
    elif health_status["overall_status"] == "degraded":
        return jsonify(health_status), 200
    else:
        return jsonify(health_status), 503


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_file_content(file_storage):
    filename = file_storage.filename
    extension = filename.rsplit('.', 1)[1].lower()
    content = ""
    file_storage.seek(0)

    if extension in ['txt', 'md']:
        content = file_storage.read().decode('utf-8')
    elif extension == 'json':
        try:
            json_data = json.load(file_storage)
            content = json.dumps(json_data, indent=2)
        except json.JSONDecodeError:
            raise ValueError("Invalid JSON file.")
    elif extension == 'pdf':
        pdf_reader = pypdf.PdfReader(file_storage)
        for page in pdf_reader.pages:
            content += page.extract_text()
    elif extension == 'docx':
        doc = docx.Document(file_storage)
        for para in doc.paragraphs:
            content += para.text + '\n'

    return content

# --- NEW: Get all documents ---


@api_bp.route('/documents', methods=['GET'])
@require_auth(["documents:read"])
def get_documents():
    """
    Fetches all documents from the database, filtered by owner_id if user is authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter documents by owner_id for authenticated user
        documents = Document.query.filter_by(
            owner_id=current_user_id).order_by(Document.created_at.desc()).all()

        results = [
            {
                "id": doc.id,
                "filename": doc.filename,
                "created_at": doc.created_at.isoformat() if doc.created_at else None
            }
            for doc in documents
        ]
        return jsonify(results)
    except Exception as e:
        print(f"An error occurred while fetching documents: {str(e)}")
        return jsonify({"error": "Failed to fetch documents"}), 500


@api_bp.route('/upload', methods=['POST'])
@require_auth(["documents:write"])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        new_document = None

        try:
            content = parse_file_content(file)

            # Get current user ID from authenticated session
            from flask import g
            current_user_id = g.user_id

            new_document = Document(
                filename=filename, content=content, owner_id=current_user_id)
            db.session.add(new_document)
            db.session.commit()

            process_and_store_document(new_document)

            # Return the new document object, matching the /documents GET route
            return jsonify({
                "message": "File uploaded and processed successfully",
                "document": {
                    "id": new_document.id,
                    "filename": new_document.filename,
                    "created_at": new_document.created_at.isoformat() if new_document.created_at else None
                }
            }), 201

        except Exception as e:
            print(f"An error occurred during file processing: {str(e)}")
            db.session.rollback()
            return jsonify({"error": f"Failed to process file: {str(e)}"}), 500

    return jsonify({"error": "File type not allowed"}), 400

# --- NEW: Delete a document ---


@api_bp.route('/documents/<int:document_id>', methods=['DELETE'])
@require_auth(["documents:write"])
def delete_document(document_id):
    """
    Deletes a document from the database and the RAG vector store, with user access validation.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Find document with user access validation
        document = Document.query.filter_by(
            id=document_id, owner_id=current_user_id).first()

        if not document:
            return jsonify({"error": "Document not found or access denied"}), 404

        # 1. Delete from RAG first.
        delete_document_from_rag(document_id)

        # 2. Delete the document object from the session.
        db.session.delete(document)

        # 3. Commit the transaction
        db.session.commit()

        return jsonify({"message": f"Document ID {document_id} and associated data deleted."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during document deletion: {str(e)}")
        return jsonify({"error": f"Failed to delete document: {str(e)}"}), 500

# --- REMOVED: /analyze ---
# This endpoint is no longer needed as the query is not user-driven.

# --- NEW: Trigger for requirements generation ---


@api_bp.route('/requirements/generate', methods=['POST'])
@require_auth(["requirements:write"])
def trigger_requirements_generation():
    """
    Triggers a full regeneration of requirements from user's documents if authenticated.
    This clears existing user requirements first.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        total_generated = generate_project_requirements(
            owner_id=current_user_id)
        return jsonify({
            "message": f"Successfully generated {total_generated} new requirements."
        })
    except Exception as e:
        print(f"An error occurred during requirements generation: {str(e)}")
        return jsonify({"error": f"Failed to generate requirements: {str(e)}"}), 500


@api_bp.route('/requirements', methods=['GET'])
@require_auth(["requirements:read"])
def get_requirements():
    """
    Fetches all requirements from the database, filtered by owner_id if user is authenticated,
    including their associated tags and the filename of their source document.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter requirements by owner_id for authenticated user
        requirements = Requirement.query.filter_by(owner_id=current_user_id).options(
            db.joinedload(Requirement.tags),
            db.joinedload(Requirement.source_document)
        ).all()

        results = []
        for req in requirements:
            results.append({
                "id": req.id,
                "req_id": req.req_id,
                "title": req.title,
                "description": req.description,
                "status": req.status,
                "priority": req.priority,
                "source_document_filename": req.source_document.filename if req.source_document else None,
                "tags": [{"id": tag.id, "name": tag.name} for tag in req.tags]
            })

        return jsonify(results)

    except Exception as e:
        print(f"An error occurred while fetching requirements: {str(e)}")
        return jsonify({"error": "Failed to fetch requirements"}), 500

# --- NEW: Get requirements count ---


@api_bp.route('/requirements/count', methods=['GET'])
@require_auth(["requirements:read"])
def get_requirements_count():
    """
    Fetches a simple count of requirements, filtered by owner_id if user is authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter count by owner_id for authenticated user
        count = db.session.query(Requirement.id).filter_by(
            owner_id=current_user_id).count()

        return jsonify({"count": count})
    except Exception as e:
        print(f"An error occurred while fetching requirements count: {str(e)}")
        return jsonify({"error": "Failed to fetch requirements count"}), 500

# --- NEW: Get project summary ---


@api_bp.route('/summary', methods=['GET'])
@require_auth(["summary:read"])
def get_summary():
    """
    Fetches the latest generated project summary from the database, filtered by owner_id if user is authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter summary by owner_id for authenticated user
        latest_summary = ProjectSummary.query.filter_by(
            owner_id=current_user_id).order_by(ProjectSummary.created_at.desc()).first()

        if latest_summary:
            return jsonify({
                "summary": latest_summary.content,
                "created_at": latest_summary.created_at.isoformat() if latest_summary.created_at else None
            })
        else:
            return jsonify({
                "summary": "No summary has been generated yet. Upload documents and click 'Regenerate Summary' on the Overview page.",
                "created_at": None
            })
    except Exception as e:
        print(f"An error occurred while fetching summary: {str(e)}")
        return jsonify({"error": "Failed to fetch summary"}), 500

# --- NEW: Trigger for summary generation ---


@api_bp.route('/summary/generate', methods=['POST'])
@require_auth(["summary:write"])
def trigger_summary_generation():
    """
    Triggers generation of a new project-wide summary.
    This retrieves context from user's documents if authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Call the summary generation function with user context
        summary_output = generate_project_summary(owner_id=current_user_id)

        # Save the new summary to the database with owner_id
        new_summary = ProjectSummary(
            content=summary_output, owner_id=current_user_id)
        db.session.add(new_summary)
        db.session.commit()

        return jsonify({
            "message": "Summary generated successfully.",
            "summary": new_summary.content,
            "created_at": new_summary.created_at.isoformat() if new_summary.created_at else None
        })
    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during summarization: {str(e)}")
        return jsonify({"error": f"Failed to summarize project: {str(e)}"}), 500


# --- Profile Management Endpoints ---

def assign_default_role_to_user(user_id, email):
    """
    Assigns a default role to a new user based on business logic.
    For now, assigns 'pilot-user' role to all new users.
    """
    try:
        # Import here to avoid circular imports
        from supertokens_python.recipe import userroles
        import asyncio

        # Determine role based on business logic
        # For now, all new users get 'pilot-user' role
        default_role = 'pilot-user'

        # Create event loop if needed
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        # Assign role to user
        async def assign_role():
            try:
                result = await userroles.add_role_to_user(
                    tenant_id="public",
                    user_id=user_id,
                    role=default_role
                )
                return result
            except Exception as e:
                print(f"Error assigning role to user {user_id}: {str(e)}")
                return None

        result = loop.run_until_complete(assign_role())

        if result:
            print(
                f"Successfully assigned role '{default_role}' to user {user_id}")
            return default_role
        else:
            print(f"Failed to assign role to user {user_id}")
            return None

    except Exception as e:
        print(f"Error in role assignment for user {user_id}: {str(e)}")
        return None


@api_bp.route('/auth/profile', methods=['POST'])
def create_profile():
    """
    Creates a new user profile.
    Expects JSON payload with user profile data.
    """
    try:
        data = request.get_json()

        # Validate required fields
        required_fields = ['user_id', 'email', 'first_name',
                           'last_name', 'company', 'job_title']
        for field in required_fields:
            if not data.get(field):
                return jsonify({"error": f"Missing required field: {field}"}), 400

        # Check if profile already exists
        existing_profile = UserProfile.query.filter_by(
            user_id=data['user_id']).first()
        if existing_profile:
            return jsonify({"error": "Profile already exists for this user"}), 409

        # Create new profile
        new_profile = UserProfile(
            user_id=data['user_id'],
            email=data['email'],
            first_name=data['first_name'],
            last_name=data['last_name'],
            company=data['company'],
            job_title=data['job_title'],
            remaining_tokens=data.get(
                'remaining_tokens', 5)  # Default to 5 tokens
        )

        db.session.add(new_profile)
        db.session.commit()

        # Assign default role to the new user
        assigned_role = assign_default_role_to_user(
            data['user_id'], data['email'])

        return jsonify({
            "message": "Profile created successfully",
            "profile": {
                "id": new_profile.id,
                "user_id": new_profile.user_id,
                "email": new_profile.email,
                "first_name": new_profile.first_name,
                "last_name": new_profile.last_name,
                "company": new_profile.company,
                "job_title": new_profile.job_title,
                "remaining_tokens": new_profile.remaining_tokens,
                "created_at": new_profile.created_at.isoformat() if new_profile.created_at else None,
                "assigned_role": assigned_role
            }
        }), 201

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during profile creation: {str(e)}")
        return jsonify({"error": f"Failed to create profile: {str(e)}"}), 500


@api_bp.route('/auth/profile', methods=['GET'])
def get_profile():
    """
    Retrieves user profile by user_id from query parameters.
    """
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({"error": "user_id parameter is required"}), 400

        profile = UserProfile.query.filter_by(user_id=user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        return jsonify({
            "profile": {
                "id": profile.id,
                "user_id": profile.user_id,
                "email": profile.email,
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "company": profile.company,
                "job_title": profile.job_title,
                "remaining_tokens": profile.remaining_tokens,
                "created_at": profile.created_at.isoformat() if profile.created_at else None,
                "updated_at": profile.updated_at.isoformat() if profile.updated_at else None
            }
        })

    except Exception as e:
        print(f"An error occurred while fetching profile: {str(e)}")
        return jsonify({"error": f"Failed to fetch profile: {str(e)}"}), 500


@api_bp.route('/profile', methods=['GET'])
@require_auth(["profile:read"])
def get_current_user_profile():
    """
    Retrieves the current authenticated user's profile.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        profile = UserProfile.query.filter_by(user_id=current_user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        # Return profile in the format expected by frontend
        return jsonify({
            "user_id": profile.user_id,
            "email": profile.email,
            "metadata": {
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "user_profile": {
                    "company": profile.company,
                    "job_title": profile.job_title
                },
                "user_token": {
                    "remaining_tokens": profile.remaining_tokens
                }
            }
        })

    except Exception as e:
        print(f"An error occurred while fetching profile: {str(e)}")
        return jsonify({"error": f"Failed to fetch profile: {str(e)}"}), 500


@api_bp.route('/profile', methods=['PUT'])
@require_auth(["profile:write"])
def update_current_user_profile():
    """
    Updates the current authenticated user's profile.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        # Extract metadata from request
        metadata = data.get('metadata', {})
        user_profile = metadata.get('user_profile', {})

        # Validate required fields
        required_fields = {
            'first_name': metadata.get('first_name'),
            'last_name': metadata.get('last_name'),
            'company': user_profile.get('company'),
            'job_title': user_profile.get('job_title')
        }

        for field, value in required_fields.items():
            if not value or str(value).strip() == '':
                return jsonify({"error": f"Missing required field: {field}"}), 400

        # Find existing profile
        profile = UserProfile.query.filter_by(user_id=current_user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        # Update profile fields
        profile.first_name = required_fields['first_name']
        profile.last_name = required_fields['last_name']
        profile.company = required_fields['company']
        profile.job_title = required_fields['job_title']

        # Update remaining tokens if provided
        user_token = metadata.get('user_token', {})
        if 'remaining_tokens' in user_token:
            profile.remaining_tokens = user_token['remaining_tokens']

        db.session.commit()

        # Return updated profile in the format expected by frontend
        return jsonify({
            "user_id": profile.user_id,
            "email": profile.email,
            "metadata": {
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "user_profile": {
                    "company": profile.company,
                    "job_title": profile.job_title
                },
                "user_token": {
                    "remaining_tokens": profile.remaining_tokens
                }
            }
        })

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred while updating profile: {str(e)}")
        return jsonify({"error": f"Failed to update profile: {str(e)}"}), 500
